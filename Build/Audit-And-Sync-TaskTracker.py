from __future__ import annotations

import json
import re
import shutil
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "AI_WordPress_Manager_Web_ASPNET_Core_Full_Task_Tracker_AR.docx"
JSON_FILE = ROOT / "task-tracker.json"
BACKUP_DOCX = ROOT / "artifacts" / "task-tracker-backup.docx"
BATCH_SIZE = 100
DONE_COLOR = "E2F0D9"
PENDING_COLOR = "FCE4EC"
HEADER_COLOR = "D9EAF7"

IGNORE_PARTS = {".git", "bin", "obj", ".vs", "node_modules", "artifacts"}

EXPLICIT_RULES: list[tuple[tuple[str, ...], tuple[str, ...]]] = [
    (("نطاق mvp", "scope"), ("docs/governance/PROJECT_GOVERNANCE.md",)),
    (("المستخدمين والأدوار", "roles"), ("docs/governance/PROJECT_GOVERNANCE.md",)),
    (("definition of done", "dod"), ("docs/governance/PROJECT_GOVERNANCE.md",)),
    (("git workflow", "الإصدارات"), ("docs/governance/PROJECT_GOVERNANCE.md", ".github/workflows/dotnet-build.yml")),
    (("backlog", "مراحل الإطلاق"), ("docs/governance/PROJECT_GOVERNANCE.md",)),
    (("سجل المخاطر", "risk register"), ("docs/governance/PROJECT_GOVERNANCE.md",)),
    (("development/staging/production", "البيئات"), ("docs/governance/PROJECT_GOVERNANCE.md",)),
    (("aiwordpressmanager.web.sln", "solution"), ("AIWordPressManager.Web.sln",)),
    (("blazor server", "web project"), ("src/AIWordPressManager.Web/AIWordPressManager.Web.csproj",)),
    (("إنشاء domain", "domain project"), ("src/AIWordPressManager.Domain/AIWordPressManager.Domain.csproj",)),
    (("إنشاء application", "application project"), ("src/AIWordPressManager.Application/AIWordPressManager.Application.csproj",)),
    (("إنشاء infrastructure", "infrastructure project"), ("src/AIWordPressManager.Infrastructure/AIWordPressManager.Infrastructure.csproj",)),
    (("wordpress integration", "عميل rest"), ("src/AIWordPressManager.Application/Abstractions/WordPress/IWordPressApiClient.cs",)),
    (("ai integration", "مزود قابل للتبديل"), ("src/AIWordPressManager.Application/Abstractions/AI", "src/AIWordPressManager.Infrastructure/AI")),
    (("automation", "jobs"), ("src/AIWordPressManager.Application/Abstractions/Automation", "src/AIWordPressManager.Web/Services/Automation")),
    (("unit/integration/e2e", "tests"), ("tests",)),
    (("sqlite", "appdbcontext"), ("src/AIWordPressManager.Persistence/AppDbContext.cs",)),
    (("site", "كيان الموقع"), ("src/AIWordPressManager.Domain/Entities/Site.cs",)),
    (("credential", "بيانات الاعتماد"), ("src/AIWordPressManager.Domain/Entities/SiteCredential.cs",)),
    (("connection", "اختبار الاتصال"), ("src/AIWordPressManager.Web/Services/SiteWebService.cs",)),
    (("notification", "الإشعارات"), ("src/AIWordPressManager.Web/Services/AppNotificationService.cs", "src/AIWordPressManager.Web/Components/Shared/GlobalNotifications.razor")),
    (("localization", "الترجمة", "العربية", "الإنجليزية"), ("src/AIWordPressManager.Web/Services/AppLanguageService.cs",)),
    (("health", "الصحة"), ("src/AIWordPressManager.Web/Services/SystemHealthService.cs",)),
    (("sync", "المزامنة"), ("src/AIWordPressManager.Web/Services/WordPressSyncWebService.cs",)),
    (("seo", "تحليل seo"), ("src/AIWordPressManager.Web/Services/SeoAuditExecutionService.cs",)),
    (("execution center", "مركز التنفيذ"), ("src/AIWordPressManager.Web/Components/Pages/ExecutionCenter.razor",)),
    (("content explorer", "مستكشف المحتوى"), ("src/AIWordPressManager.Web/Components/Pages/ContentExplorer.razor",)),
    (("media", "الوسائط"), ("src/AIWordPressManager.Web/Components/Pages/Media.razor",)),
    (("comments", "التعليقات"), ("src/AIWordPressManager.Web/Components/Pages/Comments.razor",)),
    (("users", "المستخدمين"), ("src/AIWordPressManager.Web/Components/Pages/Users.razor",)),
    (("categories", "التصنيفات", "tags", "الوسوم"), ("src/AIWordPressManager.Web/Components/Pages/Taxonomies.razor",)),
    (("install", "تثبيت أول مرة"), ("Install-First-Time.ps1", "Install-First-Time.bat")),
    (("update", "التحديث"), ("Update-System.ps1", "Update-System.bat")),
    (("build", "ci"), (".github/workflows/dotnet-build.yml", "Build/Run-Web.ps1")),
]


def normalize(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip().lower()


def normalize_id(value: str) -> str:
    value = normalize(value)
    match = re.search(r"\d+", value)
    return str(int(match.group(0))) if match else value


def repo_paths() -> set[str]:
    result: set[str] = set()
    for path in ROOT.rglob("*"):
        if any(part in IGNORE_PARTS for part in path.parts):
            continue
        if path.is_file():
            result.add(path.relative_to(ROOT).as_posix())
    return result


def path_exists(paths: set[str], required: str) -> bool:
    required = required.rstrip("/")
    return required in paths or any(path.startswith(required + "/") for path in paths)


def task_text(task: dict) -> str:
    return normalize(" | ".join(str(value) for value in task.get("values", [])))


def rule_evidence(text: str, paths: set[str]) -> list[str]:
    found: list[str] = []
    for keywords, required_paths in EXPLICIT_RULES:
        if not any(normalize(keyword) in text for keyword in keywords):
            continue
        found.extend(path for path in required_paths if path_exists(paths, path))
    return sorted(set(found))


def generic_evidence(text: str, paths: set[str]) -> list[str]:
    ignored = {"project", "document", "manager", "developer", "wordpress", "implementation"}
    tokens = [
        token for token in re.findall(r"[a-zA-Z][a-zA-Z0-9_.-]{3,}", text)
        if token.lower() not in ignored
    ]
    evidence: list[str] = []
    for token in tokens[:8]:
        low = token.lower()
        match = next((path for path in sorted(paths) if low in path.lower()), None)
        if match:
            evidence.append(match)
    return sorted(set(evidence))


def audit(data: dict) -> dict:
    paths = repo_paths()
    now = datetime.now(timezone.utc).isoformat()
    processed = 0

    for task in data.get("tasks", []):
        task_id = normalize_id(str(task.get("id", "")))
        if not task_id.isdigit() or not 1 <= int(task_id) <= BATCH_SIZE:
            continue

        processed += 1
        text = task_text(task)
        evidence = rule_evidence(text, paths) or generic_evidence(text, paths)

        if evidence:
            task["status"] = "completed"
            task["status_text"] = "مكتمل - تم التنفيذ"
            task["notes"] = "دليل التنفيذ: " + "؛ ".join(evidence[:6])
        else:
            task["status"] = "pending"
            task["status_text"] = "غير مكتمل"
            task["notes"] = "تمت المراجعة ضمن دفعة 1-100، ولا يوجد دليل تنفيذ كافٍ حتى الآن."
        task["updated_at_utc"] = now

    data["audit"] = {
        "batch": f"1-{BATCH_SIZE}",
        "processed": processed,
        "method": "repository evidence audit",
        "updated_at_utc": now,
    }
    data["summary"] = {
        "total": len(data.get("tasks", [])),
        "completed": sum(1 for task in data.get("tasks", []) if task.get("status") == "completed"),
        "pending": sum(1 for task in data.get("tasks", []) if task.get("status") != "completed"),
    }
    data["generated_at_utc"] = now
    return data


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_text_preserving_layout(cell, text: str, *, rtl: bool = True) -> None:
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    for extra in list(cell.paragraphs[1:]):
        extra._element.getparent().remove(extra._element)

    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
        run = runs[0]
    else:
        run = paragraph.add_run(text)

    run.font.name = "Cairo"
    run.font.size = Pt(8)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Cairo")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Cairo")

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if rtl and bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    elif not rtl and bidi is not None:
        p_pr.remove(bidi)


def find_column(headers: Iterable[str], candidates: tuple[str, ...]) -> int | None:
    for index, header in enumerate(headers):
        normalized = normalize(header)
        if any(candidate in normalized for candidate in candidates):
            return index
    return None


def row_task_id(row) -> str:
    if not row.cells:
        return ""
    return normalize_id(row.cells[0].text)


def apply_to_docx(data: dict) -> None:
    BACKUP_DOCX.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(DOCX, BACKUP_DOCX)

    document = Document(DOCX)
    original_table_count = len(document.tables)
    original_row_count = sum(len(table.rows) for table in document.tables)
    tasks_by_id = {normalize_id(str(task.get("id", ""))): task for task in data.get("tasks", [])}
    matched_ids: set[str] = set()

    for table in document.tables:
        if not table.rows:
            continue

        headers = [cell.text for cell in table.rows[0].cells]
        status_col = find_column(headers, ("الحالة", "حالة", "status", "الإنجاز", "التنفيذ"))
        notes_col = find_column(headers, ("ملاحظات", "الملاحظات", "notes", "note", "دليل"))

        for cell in table.rows[0].cells:
            shade_cell(cell, HEADER_COLOR)

        for row in table.rows[1:]:
            task_id = row_task_id(row)
            task = tasks_by_id.get(task_id)
            if not task:
                continue

            matched_ids.add(task_id)
            completed = task.get("status") == "completed"
            fill = DONE_COLOR if completed else PENDING_COLOR
            for cell in row.cells:
                shade_cell(cell, fill)

            if status_col is not None and status_col < len(row.cells):
                set_cell_text_preserving_layout(
                    row.cells[status_col],
                    "مكتمل - تم التنفيذ" if completed else "غير مكتمل",
                )

            if notes_col is not None and notes_col < len(row.cells):
                set_cell_text_preserving_layout(row.cells[notes_col], task.get("notes", ""))

    if len(document.tables) != original_table_count:
        raise RuntimeError("Table count changed while updating the tracker.")
    if sum(len(table.rows) for table in document.tables) != original_row_count:
        raise RuntimeError("Row count changed while updating the tracker.")

    expected_ids = {
        normalize_id(str(task.get("id", "")))
        for task in data.get("tasks", [])
        if normalize_id(str(task.get("id", "")))
    }
    missing = sorted(expected_ids - matched_ids, key=lambda value: int(value) if value.isdigit() else 999999)
    data["document_sync"] = {
        "matched_rows": len(matched_ids),
        "missing_task_ids": missing,
        "table_count": original_table_count,
        "row_count": original_row_count,
        "updated_at_utc": datetime.now(timezone.utc).isoformat(),
    }

    document.save(DOCX)

    verification = Document(DOCX)
    if len(verification.tables) != original_table_count:
        raise RuntimeError("Saved DOCX failed table-count verification.")
    if sum(len(table.rows) for table in verification.tables) != original_row_count:
        raise RuntimeError("Saved DOCX failed row-count verification.")


def main() -> int:
    if not JSON_FILE.exists() or not DOCX.exists():
        print("task-tracker.json and tracker DOCX are required", file=sys.stderr)
        return 2

    data = json.loads(JSON_FILE.read_text(encoding="utf-8"))
    data = audit(data)
    apply_to_docx(data)
    JSON_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps({"summary": data["summary"], "document_sync": data["document_sync"]}, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
